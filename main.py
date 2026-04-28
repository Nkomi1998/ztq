from fastapi.middleware.cors import CORSMiddleware
from fastapi import FastAPI, HTTPException
from database import engine, SessionLocal
from models import Base, User, UserDetails
from datetime import datetime, date
from docx import Document

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Create tables
Base.metadata.create_all(bind=engine)


# ---------------- BASIC ROOT TEST ----------------
@app.get("/")
def root():
    return {"message": "API is working"}


# ---------------- LOGIN ----------------
@app.post("/login")
def login(lastname: str, password: str):
    db = SessionLocal()

    user = db.query(User, UserDetails).join(UserDetails).filter(
        User.UserID == UserDetails.UserID,
        UserDetails.LastName == lastname,
        User.Password == password
    ).first()

    db.close()

    if not user:
        raise HTTPException(status_code=400, detail="Invalid login")

    return {"message": "Login successful"}


# ---------------- ADD USER ----------------
@app.post("/add_user")
def add_user(password: str, first_name: str, last_name: str,
             dob: str, province: str, gender: str, facilitator: bool):

    db = SessionLocal()

    new_user = User(Password=password)
    db.add(new_user)
    db.commit()
    db.refresh(new_user)

    user_details = UserDetails(
        UserID=new_user.UserID,
        FirstName=first_name,
        LastName=last_name,
        DateOfBirth=datetime.strptime(dob, "%Y-%m-%d").date(),
        Province=province,
        Gender=gender,
        Facilitator=facilitator
    )

    db.add(user_details)
    db.commit()
    db.close()

    return {"message": "User added successfully"}


# ---------------- GET USERS ----------------
@app.get("/users")
def get_users():
    db = SessionLocal()

    users = db.query(UserDetails).all()

    result = []
    for u in users:
        result.append({
            "id": u.UserID,
            "name": f"{u.FirstName} {u.LastName}",
            "dob": str(u.DateOfBirth),
            "province": u.Province,
            "gender": u.Gender,
            "facilitator": u.Facilitator
        })

    db.close()
    return result


# ---------------- EDIT USER ----------------
@app.put("/edit_user/{user_id}")
def edit_user(user_id: int, first_name: str, last_name: str):
    db = SessionLocal()

    user = db.query(UserDetails).filter(UserDetails.UserID == user_id).first()

    if not user:
        db.close()
        raise HTTPException(status_code=404, detail="User not found")

    user.FirstName = first_name
    user.LastName = last_name

    db.commit()
    db.close()

    return {"message": "User updated successfully"}


# ---------------- DELETE USER ----------------
@app.delete("/delete_user/{user_id}")
def delete_user(user_id: int):
    db = SessionLocal()

    details = db.query(UserDetails).filter(UserDetails.UserID == user_id).first()
    user = db.query(User).filter(User.UserID == user_id).first()

    if not user:
        db.close()
        raise HTTPException(status_code=404, detail="User not found")

    db.delete(details)
    db.delete(user)
    db.commit()
    db.close()

    return {"message": "User deleted successfully"}


# ---------------- EXPORT WORD ----------------
@app.get("/export_word")
def export_word():
    db = SessionLocal()

    users = db.query(UserDetails).all()

    doc = Document()
    doc.add_heading("User Report", 0)

    for u in users:
        doc.add_paragraph(
            f"{u.FirstName} {u.LastName} | {u.DateOfBirth} | {u.Province} | {u.Gender} | Facilitator: {u.Facilitator}"
        )

    file_name = "users_report.docx"
    doc.save(file_name)

    db.close()

    return {"message": f"Report saved as {file_name}"}
----For export-----
successful_download = True
if successful_download:
    return(True)
else:unsuccessful_download = False
if unsuccessful_download:
    retrun(False)